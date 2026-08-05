#!/usr/bin/env python3
"""
Insert a diagram image (PNG) into a TDD .docx at the paragraph that
contains a given placeholder string, replacing that placeholder text
with the picture, centered, scaled to fit the page width.

Use this AFTER the surrounding text of the document has already been
filled in (via the normal unzip/edit-document.xml/rezip flow described
in the docx skill). This script opens the .docx directly with
python-docx, which safely rewrites the picture relationship/media
parts for you — no manual OOXML drawing/relationship editing needed.

Usage:
    python insert_diagram.py <docx_path> <image_path> \
        [--placeholder "Insert high-level process flow diagram image here"] \
        [--width-in 6.5] \
        [--output <output_docx_path>]

If --output is omitted, the input file is overwritten in place.
Exits non-zero (and prints all paragraph text it *did* find) if the
placeholder text isn't found, so a mismatch is never silent.
"""
import argparse
import sys

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def insert_diagram(docx_path, image_path, placeholder, width_in, output_path):
    doc = Document(docx_path)

    target = None
    for p in doc.paragraphs:
        if placeholder in p.text:
            target = p
            break

    if target is None:
        print(f"ERROR: placeholder text {placeholder!r} not found in {docx_path}.")
        print("Paragraphs actually present (non-empty):")
        for p in doc.paragraphs:
            if p.text.strip():
                print(f"  - {p.text!r}")
        sys.exit(1)

    # Clear existing runs (the placeholder text) but keep the paragraph
    # itself so heading spacing / position in the doc is preserved.
    for run in list(target.runs):
        run.text = ""
        run.clear()

    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = target.add_run()
    run.add_picture(image_path, width=Inches(width_in))

    out = output_path or docx_path
    doc.save(out)
    print(f"Inserted {image_path} into {out} at paragraph matching {placeholder!r}.")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx_path")
    parser.add_argument("image_path")
    parser.add_argument(
        "--placeholder",
        default="Insert high-level process flow diagram image here",
    )
    parser.add_argument("--width-in", type=float, default=6.5)
    parser.add_argument("--output", default=None)
    args = parser.parse_args()

    insert_diagram(
        args.docx_path,
        args.image_path,
        args.placeholder,
        args.width_in,
        args.output,
    )
